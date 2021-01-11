#!/usr/bin/env python
# encoding: utf-8
# @Time    : 2021/1/11 10:00
# @Author  : lxx
# @File    : test.py
# @Software: PyCharm

import docx
from docx import Document
import xlwt;
import xlrd;
import glob
from win32com import client as wc
import os
import shutil


def readdoc(filename):
    doc = docx.Document(filename)
    tables = []
    for table in doc.tables:
        table_temp = []
        for row in table.rows:
            row_temp = []
            for cell in row.cells:
                row_temp.append(cell.text)
            table_temp.append(row_temp)
        tables.append(table_temp)
    return tables


def writeExcel(tables,filename):
    Sheet_index = 0
    workbook = xlwt.Workbook(encoding='utf-8')
    for table in tables:
        worksheet = workbook.add_sheet('sheet' + str(Sheet_index),cell_overwrite_ok = True)
        Sheet_index = Sheet_index + 1
        for rows in table:
            r = table.index(rows)
            for cell in rows:
                c = rows.index(cell)
                # print(r,c,cell)
                worksheet.write(r,c,cell)
    workbook.save(filename.split(".")[0] + ".xls")


# 将doc转换为docx
def doc2Docx(fileName,file_path,output_data_dir):
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(file_path)
    output_data_dir=output_data_dir+"/"+fileName.split(".")[0]
    if not os.path.exists(output_data_dir):
         os.mkdir(output_data_dir)
    
    doc.SaveAs(output_data_dir+"/"+fileName + "x", 12, False, "", True, "", False, False, False, False)
    doc.Close()
    # os.remove(fileName)
    word.Quit()

def get_docx_Text(filename):
    doc = docx.Document(filename)
    fullText = []
    for i in doc.paragraphs:  # 迭代docx文档里面的每一个段落
        fullText.append(i.text)  # 保存每一个段落的文本
    return '\n'.join(fullText)



if __name__ == "__main__":

    #转换所有doc为docx
    output_data_dir="D:\\LiuXianXian\\pycharm\\word_2_excel\\output_data"
    for root,dirs,files in os.walk("D:\\LiuXianXian\\pycharm\\word_2_excel\\doc_data"):
        for fileName in files:
            if fileName.endswith(".doc"):
                file_path=os.path.join(root,fileName)
                doc2Docx(fileName,file_path,output_data_dir)
            if fileName.endswith(".docx"):
                file_path=os.path.join(root,fileName)
                new_path= output_data_dir+"/"+fileName.split(".")[0]
                if not os.path.exists(new_path):
                    os.mkdir(new_path)
                shutil.copy(file_path, new_path)






    #读取docx中的表格和文字
    for root,dirs,files in os.walk("D:\\LiuXianXian\\pycharm\\word_2_excel\\output_data"):
        for fileName in files:
            if not fileName.endswith(".docx"):
                continue
            file_path = os.path.join(root,fileName)

            #读取docx中的表格保存到excel中
            tables = readdoc(file_path)
            if len(tables)>0:
                writeExcel(tables,file_path)

            #读取docx中的文本
            txt_path = file_path.split(".")[0]+".txt"
            with open(txt_path,"w",encoding="utf-8") as output:
                fullText = get_docx_Text(file_path)
                output.write(fullText)



